"""Resume parsing helpers for TalentLens uploads."""

from __future__ import annotations

import io
import json
import re
import time
from dataclasses import dataclass
from typing import Any

try:  # pragma: no cover - optional dependency
    import fitz  # type: ignore
except Exception:  # pragma: no cover
    fitz = None

try:  # pragma: no cover - optional dependency
    import pdfplumber  # type: ignore
except Exception:  # pragma: no cover
    pdfplumber = None

try:  # pragma: no cover - optional dependency
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None

try:  # pragma: no cover - optional dependency
    from PIL import Image  # type: ignore
except Exception:  # pragma: no cover
    Image = None

try:  # pragma: no cover - optional dependency
    import pytesseract  # type: ignore
except Exception:  # pragma: no cover
    pytesseract = None


SKILL_LIBRARY: list[tuple[str, tuple[str, ...]]] = [
    # ── Languages ────────────────────────────────────────────────────────────
    ("Python",               ("python",)),
    ("Java",                 ("java",)),
    ("JavaScript",           ("javascript",)),
    ("TypeScript",           ("typescript",)),
    # ── Frontend ─────────────────────────────────────────────────────────────
    ("React",                ("react",)),
    ("Node.js",              ("node.js", "nodejs")),
    ("Express",              ("express", "express.js", "expressjs")),
    ("FastAPI",              ("fastapi",)),
    ("Flask",                ("flask",)),
    ("Django",               ("django",)),
    ("HTML",                 ("html",)),
    ("CSS",                  ("css",)),
    ("Tailwind",             ("tailwind", "tailwindcss")),
    ("Vite",                 ("vite",)),
    # ── Databases ────────────────────────────────────────────────────────────
    ("SQL",                  ("sql",)),
    ("MongoDB",              ("mongodb",)),
    ("Supabase",             ("supabase",)),
    ("Firebase",             ("firebase",)),
    ("PostgreSQL",           ("postgresql", "postgres")),
    ("MySQL",                ("mysql",)),
    # ── ML / AI core ─────────────────────────────────────────────────────────
    ("Machine Learning",     ("machine learning",)),
    ("Deep Learning",        ("deep learning",)),
    ("NLP",                  ("nlp", "natural language processing")),
    ("Computer Vision",      ("computer vision",)),
    ("LLMs",                 ("llms", "llm", "large language model",
                               "large language models", "genai")),
    ("RAG",                  ("rag", "retrieval augmented generation",
                               "retrieval-augmented generation")),
    ("Sentence Transformers",("sentence transformers", "sentence-transformers")),
    ("FAISS",                ("faiss",)),
    ("LangChain",            ("langchain",)),
    ("PyTorch",              ("pytorch",)),
    ("TensorFlow",           ("tensorflow",)),
    ("Pandas",               ("pandas",)),
    ("NumPy",                ("numpy",)),
    ("Scikit-learn",         ("scikit-learn", "sklearn")),
    # ── Cloud / DevOps ───────────────────────────────────────────────────────
    ("GCP",                  ("gcp", "google cloud")),
    ("Azure",                ("azure",)),
    ("AWS",                  ("aws", "amazon web services")),
    ("Docker",               ("docker",)),
    # ── Version control — GitHub before Git so both are caught independently ─
    ("GitHub",               ("github",)),
    ("Git",                  ("git",)),
    # ── APIs / Data ──────────────────────────────────────────────────────────
    ("REST API",             ("rest api", "restful api", "api development")),
    ("Data Analysis",        ("data analysis",)),
    ("Spark",                ("spark", "apache spark")),
    ("Airflow",              ("airflow",)),
    ("ETL",                  ("etl",)),
    ("Prompt Engineering",   ("prompt engineering",)),
    ("Vector DB",            ("vector db", "vector database")),
    ("Cybersecurity",        ("cybersecurity", "cyber security")),
    ("IoT",                  ("iot", "internet of things")),
]

SECTION_HEADERS = {
    "skills",
    "education",
    "projects",
    "project",
    "experience",
    "certifications",
    "certification",
    "contact",
    "summary",
    "profile",
    "resume",
    "stream",
    "objective",
    "achievements",
    "achievement",
    "references",
    "declaration",
    "languages",
    "hobbies",
    "interests",
}

NAME_NOISE: frozenset[str] = frozenset({
    # ── Section headers ──────────────────────────────────────────────────────
    "skills", "skill", "education", "projects", "project", "experience",
    "certifications", "certification", "contact", "summary", "profile",
    "resume", "stream", "objective", "achievements", "achievement",
    "references", "declaration", "languages", "hobbies", "interests",
    # ── Generic document words that appear at the top of resumes ─────────────
    "curriculum", "vitae", "cv", "bio", "about",
    # ── Tech / tool words that look like capitalized names ([A-Z][a-z]+) ─────
    "python", "java", "javascript", "typescript", "react", "express",
    "fastapi", "flask", "django", "html", "css", "tailwind", "vite",
    "sql", "mongodb", "supabase", "firebase", "postgresql", "mysql",
    "machine", "learning", "deep", "nlp", "faiss", "langchain",
    "pytorch", "tensorflow", "pandas", "numpy", "sklearn", "scikit",
    "gcp", "azure", "aws", "docker", "git", "github", "gitlab",
    "spark", "airflow", "airlow", "analytics", "tableau", "excel",
    "kubernetes", "terraform", "prometheus", "grafana", "jenkins",
    "android", "ios", "swift", "kotlin", "golang", "rust", "scala",
    "redis", "elasticsearch", "kafka", "hadoop", "dbt", "snowflake",
    "oracle", "sqlite", "cassandra", "dynamodb",
    # ── Common single-word labels that must never be treated as a name ───────
    "developer", "engineer", "analyst", "scientist", "designer",
    "intern", "consultant", "manager", "researcher", "associate",
    "student", "undergraduate", "graduate", "trainee", "fresher",
    # ── Abbreviations that pass [A-Z] pattern ────────────────────────────────
    "ai", "ml", "rag", "nlp", "api", "ui", "ux", "hr", "cto", "ceo",
    "ms", "bs", "be", "me",
})


@dataclass
class ParsedResume:
    text: str
    candidate_profile: dict[str, Any]
    extraction_log: list[str] = None  # type: ignore[assignment]

    def __post_init__(self) -> None:
        if self.extraction_log is None:
            self.extraction_log = []


# Minimum character count for extracted text to be considered usable.
# Anything shorter is likely a scanned/image-only file or a corrupt extraction.
_MIN_RESUME_TEXT_LENGTH = 80


def debug_extraction(filename: str, payload: bytes) -> None:
    """Print raw extraction diagnostics for a resume file.

    Intended for development and debugging. Prints the raw text length
    returned by each extraction backend before any parsing occurs so you
    can see exactly what the parser will receive.

    Usage (from project root)::

        python -c "
        from backend.resume_parser import debug_extraction
        debug_extraction('resume.pdf', open('resume.pdf','rb').read())
        "
    """
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    print(f"[debug_extraction] file={filename!r}  size={len(payload)} bytes  suffix={suffix!r}")

    if suffix == "pdf":
        # PyMuPDF
        if fitz is not None:
            chunks: list[str] = []
            try:
                with fitz.open(stream=payload, filetype="pdf") as doc:
                    chunks = [page.get_text("text") for page in doc]
                raw = "\n".join(chunks)
                print(f"  PyMuPDF   : {len(raw)} chars across {len(chunks)} page(s)")
                if raw.strip():
                    print(f"  PyMuPDF preview: {raw[:120]!r}")
            except Exception as exc:
                print(f"  PyMuPDF   : ERROR — {exc}")
        else:
            print("  PyMuPDF   : not installed")

        # pdfplumber fallback
        if pdfplumber is not None:
            parts: list[str] = []
            try:
                with pdfplumber.open(io.BytesIO(payload)) as doc:
                    parts = [(page.extract_text() or "") for page in doc.pages]
                raw2 = "\n".join(parts)
                print(f"  pdfplumber: {len(raw2)} chars across {len(parts)} page(s)")
                if raw2.strip():
                    print(f"  pdfplumber preview: {raw2[:120]!r}")
            except Exception as exc:
                print(f"  pdfplumber: ERROR — {exc}")
        else:
            print("  pdfplumber: not installed")

    elif suffix == "docx":
        if Document is not None:
            try:
                doc = Document(io.BytesIO(payload))
                paragraphs = [p.text for p in doc.paragraphs]
                table_cells: list[str] = []
                for table in doc.tables:
                    for row in table.rows:
                        table_cells.append(" ".join(c.text for c in row.cells))
                raw = normalize_resume_text("\n".join(paragraphs + table_cells))
                print(f"  python-docx: {len(raw)} chars, "
                      f"{len(paragraphs)} paragraph(s), "
                      f"{len(table_cells)} table-row(s)")
                if raw.strip():
                    print(f"  python-docx preview: {raw[:120]!r}")
            except Exception as exc:
                print(f"  python-docx: ERROR — {exc}")
        else:
            print("  python-docx: not installed")

    elif suffix == "json":
        try:
            decoded = payload.decode("utf-8")
            print(f"  JSON: {len(decoded)} chars")
            print(f"  JSON preview: {decoded[:120]!r}")
        except Exception as exc:
            print(f"  JSON decode ERROR — {exc}")

    else:
        print(f"  Unsupported suffix: {suffix!r}")


def parse_uploaded_candidate(filename: str, payload: bytes) -> ParsedResume:
    """Parse an uploaded resume file and return a ``ParsedResume``.

    Extraction order
    ----------------
    PDF  → PyMuPDF first, pdfplumber fallback if PyMuPDF yields empty text.
    DOCX → python-docx.
    JSON → direct JSON parse.
    Image → pytesseract OCR.

    Raises ``ValueError`` with a clear human-readable message (never creates
    a fake candidate profile) when:
    - the extracted text is empty,
    - the extracted text is shorter than ``_MIN_RESUME_TEXT_LENGTH`` chars,
    - a required library is not installed,
    - a corrupt file causes an exception during extraction.
    """
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    log: list[str] = []

    if suffix == "json":
        profile = parse_candidate_json(payload.decode("utf-8"))
        return ParsedResume(
            text=profile.get("resumeText", ""),
            candidate_profile=profile,
            extraction_log=["JSON parsed directly from structured fields."],
        )

    if suffix == "pdf":
        text, log = _extract_pdf_text_with_log(payload)
        _assert_usable_text(text, "PDF", log)
        return ParsedResume(text=text, candidate_profile=parse_candidate_text(text),
                            extraction_log=log)

    if suffix == "docx":
        text, log = _extract_docx_text_with_log(payload)
        _assert_usable_text(text, "DOCX", log)
        return ParsedResume(text=text, candidate_profile=parse_candidate_text(text),
                            extraction_log=log)

    if suffix in {"png", "jpg", "jpeg", "webp"}:
        text = extract_image_text(payload)
        log = [f"Image OCR extracted {len(text)} chars."]
        if not text.strip():
            raise ValueError(
                "Unable to extract text from this image. "
                "Ensure the image is clear and contains readable text, "
                "or upload a PDF/DOCX/JSON instead."
            )
        _assert_usable_text(text, "image", log)
        return ParsedResume(text=text, candidate_profile=parse_candidate_text(text),
                            extraction_log=log)

    raise ValueError(
        f"Unsupported file type '.{suffix}'. "
        "Please upload a PDF, DOCX, or JSON resume."
    )


def _extract_pdf_text_with_log(payload: bytes) -> tuple[str, list[str]]:
    """Extract text from a PDF payload, trying PyMuPDF then pdfplumber.

    Returns ``(text, log)`` where *log* records which backend was used and
    how many characters were extracted. Never raises — extraction errors are
    recorded in *log* and an empty string is returned so the caller can
    raise a clean ``ValueError`` via ``_assert_usable_text``.
    """
    log: list[str] = []

    # ── Pass 1: PyMuPDF ───────────────────────────────────────────────────────
    if fitz is not None:
        try:
            chunks: list[str] = []
            with fitz.open(stream=payload, filetype="pdf") as doc:
                for page in doc:
                    chunks.append(page.get_text("text"))
            text = normalize_resume_text("\n".join(chunks))
            log.append(
                f"PyMuPDF: extracted {len(text)} chars from {len(chunks)} page(s)."
            )
            if text.strip():
                return text, log
            log.append("PyMuPDF returned empty text — trying pdfplumber fallback.")
        except Exception as exc:
            log.append(f"PyMuPDF error: {exc} — trying pdfplumber fallback.")
    else:
        log.append("PyMuPDF (fitz) not installed — trying pdfplumber.")

    # ── Pass 2: pdfplumber fallback ───────────────────────────────────────────
    if pdfplumber is not None:
        try:
            parts: list[str] = []
            with pdfplumber.open(io.BytesIO(payload)) as doc:
                for page in doc.pages:
                    parts.append(page.extract_text() or "")
            text = normalize_resume_text("\n".join(parts))
            log.append(
                f"pdfplumber: extracted {len(text)} chars from {len(parts)} page(s)."
            )
            return text, log
        except Exception as exc:
            log.append(f"pdfplumber error: {exc}")
    else:
        log.append("pdfplumber not installed.")

    log.append("All PDF extraction backends failed or returned empty text.")
    return "", log


def _extract_docx_text_with_log(payload: bytes) -> tuple[str, list[str]]:
    """Extract text from a DOCX payload using python-docx.

    Returns ``(text, log)``. Handles corrupt files gracefully — errors are
    recorded in *log* rather than propagating as unhandled exceptions.
    """
    log: list[str] = []

    if Document is None:
        log.append(
            "python-docx is not installed. "
            "Install it with: pip install python-docx"
        )
        return "", log

    try:
        doc = Document(io.BytesIO(payload))
        paragraphs = [p.text for p in doc.paragraphs]
        table_cells: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                table_cells.append(" ".join(cell.text for cell in row.cells))
        text = normalize_resume_text("\n".join(paragraphs + table_cells))
        log.append(
            f"python-docx: extracted {len(text)} chars from "
            f"{len(paragraphs)} paragraph(s) and "
            f"{len(table_cells)} table-row(s)."
        )
        return text, log
    except Exception as exc:
        log.append(f"python-docx error: {exc}")
        return "", log


def _assert_usable_text(text: str, file_type: str, log: list[str]) -> None:
    """Raise ``ValueError`` when extracted *text* is too short to be useful.

    This is the single gate that prevents fake candidate profiles from being
    created from empty or near-empty extractions (e.g. scanned image PDFs,
    password-protected files, corrupt uploads).
    """
    char_count = len(text.strip())
    if char_count == 0:
        detail = " | ".join(log) if log else "no extraction log available"
        raise ValueError(
            f"Unable to extract any text from the uploaded {file_type}. "
            "The file may be scanned/image-only, password-protected, or corrupt. "
            "Please upload a text-based PDF, DOCX, or JSON instead. "
            f"[extraction log: {detail}]"
        )
    if char_count < _MIN_RESUME_TEXT_LENGTH:
        detail = " | ".join(log) if log else "no extraction log available"
        raise ValueError(
            f"Extracted text from {file_type} is too short to parse reliably "
            f"({char_count} chars, minimum {_MIN_RESUME_TEXT_LENGTH}). "
            "The file may be partially scanned, corrupt, or nearly empty. "
            "Please upload a clearer PDF/DOCX/JSON. "
            f"[extraction log: {detail}]"
        )


# ── Kept as public helpers for backwards compatibility and direct callers ─────

def extract_pdf_text(payload: bytes) -> str:
    """Public wrapper — returns extracted text or empty string.

    Prefer ``_extract_pdf_text_with_log`` internally so callers get the
    extraction log. This function is retained for external callers and tests.
    """
    text, _ = _extract_pdf_text_with_log(payload)
    return text


def extract_docx_text(payload: bytes) -> str:
    """Public wrapper — returns extracted text or empty string."""
    text, _ = _extract_docx_text_with_log(payload)
    return text


def extract_image_text(payload: bytes) -> str:
    if Image is None or pytesseract is None:
        return ""
    image = Image.open(io.BytesIO(payload))
    return normalize_resume_text(pytesseract.image_to_string(image))


def parse_candidate_json(text: str) -> dict[str, Any]:
    parsed = json.loads(text)
    if isinstance(parsed, list):
        parsed = next((item for item in parsed if isinstance(item, dict)), None)
        if parsed is None:
            raise ValueError("Uploaded JSON must contain at least one candidate object.")
    if not isinstance(parsed, dict):
        raise ValueError("Uploaded JSON must be a candidate object.")
    profile = parsed.get("profile") if isinstance(parsed, dict) and isinstance(parsed.get("profile"), dict) else parsed
    skills = normalize_skills((parsed.get("skills") if isinstance(parsed, dict) else None) or profile.get("skills") or [])
    projects = normalize_list((parsed.get("projects") if isinstance(parsed, dict) else None) or profile.get("projects") or [])
    education_raw = (parsed.get("education") if isinstance(parsed, dict) else None) or profile.get("education") or []
    education_list = normalize_list(education_raw)
    degree = normalize_text(profile.get("degree") or (parsed.get("degree") if isinstance(parsed, dict) else ""))
    college = normalize_text(profile.get("college") or (parsed.get("college") if isinstance(parsed, dict) else ""))
    role = clean_role(profile.get("headline") or profile.get("role") or infer_role_from_skills(skills))
    experience = int(float(profile.get("years_experience") or parsed.get("years_experience") or 0)) if isinstance(parsed, dict) else 0
    education = format_education(degree, college, " - ".join(education_list))
    experience_summary = build_experience_summary("", projects, experience, role, degree)
    profile_label = build_profile_label(role, degree, experience_summary)
    candidate_name = pick_name_from_value(
        profile.get("name") or profile.get("full_name") or (parsed.get("name") if isinstance(parsed, dict) else "")
    ) or "Candidate Profile"
    display_role = profile_label if (profile_label and not _is_forbidden_label(profile_label)) else role
    confidence = compute_confidence(
        candidate_name, skills, projects, degree, college,
        bool(profile.get("email")), bool(profile.get("phone")),
        bool(profile.get("linkedin") or profile.get("linkedin_url")),
        bool(profile.get("github") or profile.get("github_url")),
    )
    return {
        "candidate_id": f"USER_{int(time.time() * 1000)}",
        "name": candidate_name,
        "candidate_name": candidate_name,
        "profile_label": profile_label,
        "role": display_role,
        "education": education,
        "college": college,
        "degree": degree,
        "skills": skills,
        "projects": projects,
        "experience": experience,
        "experience_summary": experience_summary,
        "email": normalize_text(profile.get("email") or (parsed.get("email") if isinstance(parsed, dict) else "")),
        "phone": normalize_text(profile.get("phone") or (parsed.get("phone") if isinstance(parsed, dict) else "")),
        "linkedin": normalize_text(profile.get("linkedin_url") or profile.get("linkedin") or (parsed.get("linkedin") if isinstance(parsed, dict) else "")),
        "github": normalize_text(profile.get("github_url") or profile.get("github") or (parsed.get("github") if isinstance(parsed, dict) else "")),
        "source": "USER_ADDED",
        "parsing_confidence": confidence,
        "resumeText": normalize_resume_text(text),
    }


def parse_candidate_text(text: str) -> dict[str, Any]:
    cleaned = normalize_resume_text(text)
    lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
    skills = extract_skills(cleaned)
    projects = extract_section_items(cleaned, "projects?|project experience")
    experience = infer_experience_years(cleaned, projects)
    degree = infer_degree(cleaned)
    college = infer_college(cleaned)
    role = infer_role(cleaned, skills)
    education = format_education(degree, college, "")
    candidate_name = pick_candidate_name(lines, cleaned) or "Candidate Profile"
    email_match = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", cleaned, re.I)
    phone_match = re.search(r"(?:\+?91[-\s]?)?[6-9]\d{9}", cleaned)
    linkedin_match = re.search(r"(?:https?://)?(?:www\.)?linkedin\.com/[^\s)]+", cleaned, re.I)
    github_match = re.search(r"(?:https?://)?(?:www\.)?github\.com/[^\s)]+", cleaned, re.I)
    experience_summary = build_experience_summary(cleaned, projects, experience, role, degree)
    profile_label = build_profile_label(role, degree, experience_summary)

    # The `role` field shown in cards must never be a placeholder.
    # Use profile_label as the display role when it's more meaningful.
    display_role = profile_label if (profile_label and not _is_forbidden_label(profile_label)) else role

    confidence = compute_confidence(
        candidate_name, skills, projects, degree, college,
        bool(email_match), bool(phone_match),
        bool(linkedin_match), bool(github_match),
    )
    return {
        "candidate_id": f"USER_{int(time.time() * 1000)}",
        "name": candidate_name,
        "candidate_name": candidate_name,
        "profile_label": profile_label,
        "role": display_role,
        "education": education or "Candidate Profile",
        "college": college,
        "degree": degree,
        "skills": skills,
        "projects": projects,
        "experience": experience,
        "experience_summary": experience_summary,
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0) if phone_match else "",
        "linkedin": linkedin_match.group(0) if linkedin_match else "",
        "github": github_match.group(0) if github_match else "",
        "source": "USER_ADDED",
        "parsing_confidence": confidence,
        "resumeText": cleaned,
    }


def normalize_resume_text(text: str) -> str:
    cleaned = text.replace("\r", "\n").replace("\u00a0", " ")
    cleaned = re.sub(r"[•●▪◦·]", "\n", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def normalize_text(value: Any) -> str:
    return str(value or "").strip()


def normalize_list(value: Any) -> list[str]:
    if isinstance(value, list):
        values: list[str] = []
        for item in value:
            if isinstance(item, dict):
                text = item.get("name") or item.get("title") or item.get("degree") or item.get("school") or item.get("description")
            else:
                text = item
            if normalize_text(text):
                values.append(normalize_text(text))
        return values
    if not value:
        return []
    return [item.strip() for item in re.split(r"[\n,;|]", str(value)) if item.strip()]


def normalize_skills(value: Any) -> list[str]:
    text = " ".join(normalize_list(value)) if isinstance(value, list) else str(value or "")
    return extract_skills(text)


def clean_name(value: Any) -> str:
    """Strip non-name characters and collapse whitespace."""
    return re.sub(r"\s+", " ", re.sub(r"[^A-Za-z .'-]", " ", str(value or ""))).strip()


def pick_name_from_value(value: Any) -> str:
    """Validate and return a name string from a structured field value.

    Used when the name comes from a JSON ``name`` field rather than free text.
    Applies the same token rules as ``pick_candidate_name`` to reject
    tech words, section headers, and single-token values.
    """
    clean = clean_name(value)
    tokens = clean.split()
    if (
        2 <= len(tokens) <= 4
        and all(is_name_token(t) for t in tokens)
        and any(len(t.rstrip(".")) > 1 for t in tokens)
    ):
        return clean
    return ""


def is_name_token(token: str) -> bool:
    """Return True when *token* could be part of a human name.

    Acceptance criteria (all must hold):
    - After stripping a trailing dot (initial like "M."), the remaining string
      matches ``[A-Z][a-z]+`` (capitalized word) or ``[A-Z]`` (single initial).
    - The lowercase form must NOT be in NAME_NOISE (blocks tech/section words).

    Examples that pass : "Shirley", "S", "Nair", "Kumar", "M"
    Examples that fail  : "Python", "Docker", "Azure", "AI", "ML", "GitHub"
    """
    # Strip a single trailing period (handles initials like "M.")
    stripped = token.rstrip(".")
    if not stripped:
        return False
    if stripped.lower() in NAME_NOISE:
        return False
    return bool(re.fullmatch(r"[A-Z][a-z]+|[A-Z]", stripped))


def pick_candidate_name(lines: list[str], text: str) -> str:
    """Extract the candidate name from the top of a resume.

    Strategy
    --------
    1. Labeled line first: ``Name: Shirley S`` / ``Full Name: Priya Nair``
       (case-insensitive, colon or dash separator).
    2. Scan the first 5 non-empty lines for a plausible name:
       - Strip inline contact noise (email, phone, linkedin, github, cgpa).
       - Strip degree keywords that sometimes follow the name on the same line.
       - Accept 2–4 tokens where every token passes ``is_name_token``.
       - Require at least one full word (not just an initial) for confidence.
       - Skip lines containing digits, @ signs, or URLs.
       - Skip recognised section headers.
    3. Return ``""`` (caller substitutes "Candidate Profile") if nothing found.

    Never uses the filename or any placeholder string as the name.
    """
    # ── Pass 1: labeled "Name:" line anywhere in the resume ──────────────────
    labeled = re.search(
        r"(?:candidate\s+name|full\s+name|name)\s*[:\-]\s*([A-Za-z][A-Za-z .''-]{1,60})",
        text,
        re.I,
    )
    if labeled:
        raw = labeled.group(1).split("\n")[0].strip()
        tokens = clean_name(raw).split()
        if (
            2 <= len(tokens) <= 4
            and all(is_name_token(t) for t in tokens)
            and any(len(t.rstrip(".")) > 1 for t in tokens)
        ):
            return clean_name(raw)

    # ── Pass 2: first 5 non-empty lines ──────────────────────────────────────
    for line in lines[:5]:
        # Remove inline contact fields that may appear on the same line
        clean = re.sub(
            r"\b(email|phone|mobile|tel|linkedin|github|cgpa|gpa|percentage)\b.*$",
            "",
            line,
            flags=re.I,
        )
        # Remove degree keywords that sometimes appear directly after the name
        clean = re.sub(
            r"\b(b\.?\s?tech|b\.?\s?e\.?|m\.?\s?tech|m\.?\s?e\.?|msc|bsc|bca|mca"
            r"|bachelor|master|phd)\b.*$",
            "",
            clean,
            flags=re.I,
        )
        clean = clean_name(clean)

        # Hard rejections
        if not clean or len(clean) > 50:
            continue
        if re.search(r"\d|@|https?://|www\.", line):
            continue
        if is_section_header(clean):
            continue

        tokens = clean.split()
        if len(tokens) < 2 or len(tokens) > 4:
            continue

        # All tokens must look like name tokens
        if not all(is_name_token(t) for t in tokens):
            continue

        # At least one full word (not just an initial) required
        if not any(len(t.rstrip(".")) > 1 for t in tokens):
            continue

        return clean

    return ""


def is_section_header(text: str) -> bool:
    return normalize_text(text).strip(" :-").lower() in SECTION_HEADERS


def extract_skills(text: str) -> list[str]:
    """Extract canonical skills from *text* using the controlled skill library.

    Design
    ------
    - Each alias is matched as a whole-word / whole-phrase using ``\\b`` word
      boundaries so that ``"rag"`` matches ``"RAG"`` but not ``"fragrant"``,
      and ``"llm"`` matches ``"LLM"`` but not ``"llms"`` unless that alias is
      also listed.
    - Matching is case-insensitive.
    - Multi-word aliases (e.g. ``"rest api"``) match anywhere in the text.
    - Short all-caps tokens (``"sql"``, ``"rag"``, ``"llm"``, ``"git"``) are
      common in skill lists separated by commas, newlines, or spaces — word
      boundaries handle all of those correctly without fragile space padding.
    - GitHub is listed before Git so both are detected independently;
      ``"github"`` does *not* trigger ``"git"`` because ``\\bgit\\b`` won't
      match inside ``"github"``.
    - Duplicates are impossible because each canonical name appears once.
    """
    lowered = text.lower()
    found: list[str] = []
    for canonical, aliases in SKILL_LIBRARY:
        for alias in aliases:
            # Build a word-boundary regex. Multi-word aliases (contain spaces
            # or hyphens) use boundaries only at the outer edges.
            pattern = r"\b" + re.escape(alias) + r"\b"
            if re.search(pattern, lowered):
                found.append(canonical)
                break  # canonical matched — no need to check remaining aliases
    return found


def extract_section_items(text: str, labels: str) -> list[str]:
    match = re.search(
        rf"(?:{labels})\s*[:\-]?\s*([\s\S]{{0,900}}?)(?=\n\s*(?:education|skills|experience|certifications?|projects?|summary|objective|achievements?|contact)\b|$)",
        text,
        re.I,
    )
    if not match:
        return []
    values = [item.strip(" -*") for item in re.split(r"[\n;]", match.group(1)) if item.strip()]
    return [item for item in values if len(item) > 3 and not is_section_header(item)][:6]


def infer_experience_years(text: str, projects: list[str]) -> int:
    explicit = re.search(r"(\d+(?:\.\d+)?)\+?\s*(?:years|yrs|year)\b", text, re.I)
    if explicit:
        return max(0, min(20, round(float(explicit.group(1)))))
    current_year = 2026
    month_map = {"jan": 0, "feb": 1, "mar": 2, "apr": 3, "may": 4, "jun": 5, "jul": 6, "aug": 7, "sep": 8, "oct": 9, "nov": 10, "dec": 11}
    months = 0
    range_pattern = re.compile(
        r"\b(?:(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*)?(20\d{2}|19\d{2})\s*(?:-|to)\s*(?:(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s*)?(present|current|20\d{2}|19\d{2})\b",
        re.I,
    )
    for match in range_pattern.finditer(text):
        start_month = month_map.get((match.group(1) or "jan")[:3].lower(), 0)
        start_year = int(match.group(2))
        if re.fullmatch(r"present|current", match.group(4), re.I):
            end_year = current_year
            end_month = 5
        else:
            end_year = int(match.group(4))
            end_month = month_map.get((match.group(3) or "dec")[:3].lower(), 11)
        months += max(1, (end_year - start_year) * 12 + (end_month - start_month))
    if months > 0:
        return max(1, min(20, round(months / 12)))
    years = [int(match.group(0)) for match in re.finditer(r"\b(?:19|20)\d{2}\b", text)]
    if len(years) >= 2:
        return max(1, min(12, max(years) - min(years)))
    internships = len(re.findall(r"\b(intern|internship|trainee|apprentice)\b", text, re.I))
    work_signals = len(re.findall(r"\b(experience|worked|developed|built|engineer|developer|analyst|freelance|startup|consultant)\b", text, re.I))
    project_signals = len(re.findall(r"\b(built|developed|designed|deployed|implemented|created|launched|maintained)\b", " ".join(projects), re.I))
    if internships >= 2 or work_signals >= 8 or project_signals >= 3:
        return 2
    if internships or work_signals >= 3 or project_signals >= 1:
        return 1
    return 0


def infer_degree(text: str) -> str:
    """Extract degree string including specialisation from resume text."""
    patterns = [
        # Full spelled-out degree with possible specialisation
        r"\b(Bachelor of Technology|Master of Technology|Bachelor of Engineering|Master of Engineering|Bachelor of Science|Master of Science)\b[^\n,;]{0,60}",
        # Abbreviated form: B.Tech/M.Tech/B.E etc. with optional specialisation
        r"\b(B\.?\s?Tech|B\.?\s?E\.?|M\.?\s?Tech|M\.?\s?E\.?|MSc|M\.?\s?Sc|BSc|B\.?\s?Sc|MBA|BCA|MCA|PhD|MS|BE|ME)\b[^\n,;]{0,60}",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.I)
        if match:
            raw = match.group(0)
            # Remove trailing noise (email, URL, CGPA)
            raw = re.split(r"\s{2,}|\s*[|,]\s*", raw)[0]
            result = clean_role(raw)
            if result and len(result) > 2:
                return result
    return ""


def infer_college(text: str) -> str:
    """Extract institution name from resume text.

    Prefers lines that contain a known institution keyword. Falls back to
    context-based extraction (at/from/of).
    """
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        if re.search(r"(university|college|institute|school|iit\b|nit\b|iiit\b|bits\b|karunya|anna\s+university|psg|pes\s+university)", line, re.I):
            if not is_section_header(line):
                # Remove any leading bullet or dash
                cleaned = re.sub(r"^[-*•·]\s*", "", line).strip()
                # Strip trailing contact-like content
                cleaned = re.split(r"[|,;]", cleaned)[0].strip()
                result = clean_role(cleaned)
                if result and len(result) > 3:
                    return result
    # Fallback: look for "at/from/of <Institution>"
    match = re.search(
        r"(?:at|from|of)\s+([A-Z][A-Za-z&.,'()\- ]{2,80}(?:University|College|Institute|School|Campus|Sciences|Technology))",
        text,
    )
    return clean_role(match.group(1)) if match else ""


def infer_role(text: str, skills: list[str]) -> str:
    """Infer a role/headline from resume text."""
    # Check for an explicit role/headline line in the first 8 lines
    headline = next(
        (
            line.strip()
            for line in text.splitlines()[:8]
            if re.search(r"\b(engineer|developer|analyst|scientist|designer|intern|consultant|manager|researcher|associate)\b", line, re.I)
            and not is_section_header(line.strip())
        ),
        "",
    )
    if headline:
        return clean_role(headline)
    lowered = text.lower()
    if re.search(r"\b(student|undergraduate|graduate)\b", lowered):
        return "Student / Intern Profile"
    return infer_role_from_skills(skills)


def infer_role_from_skills(skills: list[str]) -> str:
    lowered = {skill.lower() for skill in skills}
    if {"machine learning", "deep learning"} & lowered or {"pytorch", "rag", "llms", "sentence transformers"} & lowered:
        return "AI/ML Engineer"
    if "sql" in lowered and ("spark" in lowered or "airflow" in lowered or "etl" in lowered):
        return "Data Engineer"
    if {"react", "javascript", "node.js", "typescript"} & lowered:
        return "Software Engineer"
    if {"sql", "data analysis", "pandas"} & lowered:
        return "Data Scientist"
    return "Candidate Profile"


def clean_role(value: Any) -> str:
    cleaned = re.sub(r"\s+", " ", re.sub(r"[^A-Za-z /&+.-]", " ", str(value or ""))).strip()
    cleaned = re.sub(r"\b(resume|curriculum|vitae|profile|summary|objective)\b", "", cleaned, flags=re.I).strip()
    return cleaned[:60] if cleaned else ""


def degree_short_label(degree: str) -> str:
    """Convert a full degree string to a short label like 'B.Tech CSE' or 'B.E CSE'.

    Priority order for prefix detection is most-specific first so that
    'B.Tech' is not accidentally matched as 'B.E' or vice-versa.
    """
    if not degree:
        return ""
    lowered = degree.lower().strip()

    # ── Prefix detection (most-specific patterns first) ──────────────────────
    if re.search(r"bachelor of technology|\bb\.?\s?tech\b|b tech\b", lowered):
        prefix = "B.Tech"
    elif re.search(r"master of technology|\bm\.?\s?tech\b|m tech\b", lowered):
        prefix = "M.Tech"
    elif re.search(r"bachelor of engineering|\bb\.?\s?e\.?\s(?:cse|cs|it|ai|ds|ece|eee|mech|\b)|^b\.?e\.?$", lowered):
        prefix = "B.E"
    elif re.search(r"master of engineering|\bm\.?\s?e\.?\s(?:cse|cs|it|ai|\b)|^m\.?e\.?$", lowered):
        prefix = "M.E"
    elif re.search(r"\bm\.?\s?sc\b|^msc\b", lowered):
        prefix = "M.Sc"
    elif re.search(r"\bb\.?\s?sc\b|^bsc\b", lowered):
        prefix = "B.Sc"
    elif re.search(r"\bmba\b", lowered):
        prefix = "MBA"
    elif re.search(r"\bphd\b|\bp\.?h\.?d\.?\b", lowered):
        prefix = "PhD"
    elif re.search(r"\bbca\b", lowered):
        prefix = "BCA"
    elif re.search(r"\bmca\b", lowered):
        prefix = "MCA"
    else:
        return degree

    # ── Specialisation suffix detection ──────────────────────────────────────
    if re.search(r"computer science and engineering|\bcse\b", lowered):
        suffix = " CSE"
    elif re.search(r"computer science\b", lowered):
        suffix = " CS"
    elif re.search(r"information technology|\bit\b", lowered):
        suffix = " IT"
    elif re.search(r"artificial intelligence\b", lowered):
        suffix = " AI"
    elif re.search(r"data science\b", lowered):
        suffix = " DS"
    elif re.search(r"electronics and communication|\bece\b", lowered):
        suffix = " ECE"
    elif re.search(r"electrical engineering|\beee\b", lowered):
        suffix = " EEE"
    elif re.search(r"mechanical engineering|\bmech\b", lowered):
        suffix = " Mech"
    else:
        suffix = ""
    return f"{prefix}{suffix}".strip()


def format_education(degree: str, college: str, fallback: str) -> str:
    pieces = [piece for piece in (degree, college) if piece]
    if pieces:
        return " - ".join(pieces)
    return fallback if fallback and fallback.lower() not in {"not specified", "stream"} else ""


def build_experience_summary(text: str, projects: list[str], experience: int, role: str, degree: str = "") -> str:
    """Build a human-readable experience summary string.

    Rules:
    - experience > 1 full years → "{n} years experience"
    - internship signal + student degree/words → "Student / Intern Profile"
    - internship signal alone (no student degree) → "Internship Experience"
    - student degree/words + no experience → "Student / Intern Profile"
    - experience == 1 year → "1 year experience"  (only if no internship/student signal)
    - project evidence but no numeric years → "Experience evidence available"
    - otherwise → "Candidate Profile"
    """
    combined = f"{role} {text} {degree}".lower()
    degree_lower = degree.lower()

    is_student_degree = bool(re.search(
        r"(b\.tech|btech|bachelor\s+of\s+technology|bachelor\s+of\s+engineering|b\.e\b)", degree_lower
    ))
    has_student_words = bool(re.search(r"\b(student|undergraduate|graduate)\b", combined))
    has_internship = bool(re.search(r"\b(intern|internship|trainee|apprentice)\b", combined))

    # Rule 1: Clear professional experience (> 1 year, not intern/student context)
    if experience > 1 and not (has_student_words or is_student_degree):
        return f"{experience} year{'s' if experience != 1 else ''} experience"

    # Rule 2: Internship + student context
    if has_internship:
        if has_student_words or is_student_degree:
            return "Student / Intern Profile"
        return "Internship Experience"

    # Rule 3: Pure student profile (degree + no work experience)
    if has_student_words or (is_student_degree and experience == 0):
        return "Student / Intern Profile"

    # Rule 4: B.Tech student with some experience years counted (project builds etc.)
    if is_student_degree and experience <= 1:
        return "Student / Intern Profile"

    # Rule 5: Numeric experience on a non-student profile
    if experience > 0:
        return f"{experience} year{'s' if experience != 1 else ''} experience"

    # Rule 6: Project/work signals without clear years
    if re.search(r"\b(worked|developed|built|freelance|consultant|startup|maintained|deployed)\b", combined) or projects:
        return "Experience evidence available"

    return "Candidate Profile"


_FORBIDDEN_LABELS = frozenset({
    "imported candidate",
    "data candidate",
    "uploaded candidate",
    "resume profile",
    "resume-based talent profile",
    "not specified",
    "stream",
    "resume",
    "filename",
    "",
})


def _is_forbidden_label(value: str) -> bool:
    """Return True if value is a placeholder/forbidden display label."""
    return value.strip().lower() in _FORBIDDEN_LABELS


def build_profile_label(role: str, degree: str, experience_summary: str) -> str:
    """Build a human-readable, recruiter-facing profile label.

    Rules (in priority order):
    1. Student or intern profile → "<Degree Short Label> Student"
       e.g. "B.Tech CSE Student", "B.E CSE Student"
    2. Internship-only evidence without student degree → "Internship Experience"
    3. Explicit non-placeholder role → role as-is
    4. Degree short label alone → e.g. "B.Tech CSE"
    5. Fallback → "Candidate Profile"

    Never returns a forbidden placeholder string.
    """
    degree_label = degree_short_label(degree)

    if experience_summary == "Student / Intern Profile":
        label = f"{degree_label} Student".strip() if degree_label else "Student / Intern Profile"
        return label

    if experience_summary == "Internship Experience":
        # Keep role if it's meaningful (e.g. "Software Engineering Intern")
        if role and not _is_forbidden_label(role) and role != "Student / Intern Profile":
            return role
        return degree_label or "Internship Experience"

    # Use the explicitly inferred role when it is meaningful
    if role and not _is_forbidden_label(role) and role not in ("Candidate Profile", "Student / Intern Profile"):
        return role

    return degree_label or "Candidate Profile"


def compute_confidence(
    name: str,
    skills: list[str],
    projects: list[str],
    degree: str,
    college: str,
    has_email: bool,
    has_phone: bool,
    has_linkedin: bool = False,
    has_github: bool = False,
) -> int:
    """Compute parsing confidence score 0–100."""
    score = 0
    if name and name != "Candidate Profile":
        score += 26
    if has_email:
        score += 10
    if has_phone:
        score += 8
    if has_linkedin:
        score += 6
    if has_github:
        score += 6
    if degree:
        score += 14
    if college:
        score += 12
    score += min(20, len(skills) * 3)
    score += min(10, len(projects) * 3)
    return min(100, score)
